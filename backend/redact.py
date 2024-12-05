import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from PIL import Image
from PIL.ExifTags import TAGS
import cv2
import numpy as np
import fitz  # PyMuPDF
from transformers import pipeline
import pytesseract
from PIL import ImageDraw
import io
import speech_recognition as sr
from pydub import AudioSegment
AudioSegment.converter = "./ffmpeg/bin/ffmpeg.exe"
from pydub.generators import WhiteNoise
import time

# Load the NER pipeline for detecting sensitive data
ner_pipeline = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english", grouped_entities=True)

# Existing redaction functions (redact_text_in_file, redact_pdf, redact_image, redact_faces, etc.)

def redact_audio(audio_path, output_filename):
    try:
        recognizer = sr.Recognizer()
        audio = AudioSegment.from_file(audio_path)
        temp_wav_path = "./temp.wav"
        
        # Convert to WAV format if necessary
        audio.export(temp_wav_path, format="wav")

        # Load and transcribe audio
        with sr.AudioFile(temp_wav_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)

        # Perform NER on transcribed text
        entities = ner_pipeline(text)
        sensitive_words = [entity['word'] for entity in entities if entity['entity_group'] in ["PER", "LOC", "ORG"]]

        # Replace sensitive words with white noise in the audio
        redacted_audio = audio
        for word in sensitive_words:
            start = text.find(word) / len(text) * len(audio)  # Approximate position
            duration = len(word) / len(text) * len(audio)
            redacted_audio = redacted_audio.overlay(WhiteNoise().to_audio_segment(duration=duration), position=start)

        # Save the redacted audio
        redacted_audio_path = f"./redacted/{output_filename}"
        redacted_audio.export(redacted_audio_path, format="wav")

        return output_filename
    except Exception as e:
        print(f"Error redacting audio: {e}")
        return None

def redact_text_in_file(file_path, output_filename):
    try:
        _, ext = os.path.splitext(file_path)

        if ext.lower() == '.pdf':
            return redact_pdf(file_path, output_filename)
        elif ext.lower() in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
            return redact_image(file_path, output_filename)
        else:
            print("Unsupported file type. Please upload a PDF or image file.")
            return None

    except Exception as e:
        print(f"Error processing file: {e}")
        return None

def redact_pdf(pdf_path, output_filename):
    try:
        doc = fitz.open(pdf_path)
        redacted_pdf_path = f"./redacted/{output_filename}"

        for page_num in range(doc.page_count):
            page = doc[page_num]

            # Step 1: Redact sensitive text directly from the PDF text
            page_text = page.get_text("text")
            entities = ner_pipeline(page_text)
            for entity in entities:
                if entity['entity_group'] in ["PER", "LOC", "ORG", "MISC"]:
                    sensitive_text = entity['word']
                    for instance in page.search_for(sensitive_text):
                        rect = fitz.Rect(instance)
                        page.draw_rect(rect, color=(0, 0, 0), fill=(0, 0, 0))

            # Step 2: Redact sensitive information within images on the page
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = doc.extract_image(xref)
                img_data = base_image["image"]

                # Convert image data to PIL Image for editing
                pil_img = Image.open(io.BytesIO(img_data))
                ocr_data = pytesseract.image_to_data(pil_img, output_type=pytesseract.Output.DICT)

                # Draw over sensitive text detected in the image
                draw = ImageDraw.Draw(pil_img)
                for i, word in enumerate(ocr_data['text']):
                    if word.strip():  # Ensure word is not empty
                        text_entities = ner_pipeline(word)
                        for entity in text_entities:
                            if entity['entity_group'] in ["PER", "LOC", "ORG", "MISC"]:
                                # Get bounding box and draw a black rectangle over the sensitive word
                                x, y, w, h = (ocr_data['left'][i], ocr_data['top'][i],
                                              ocr_data['width'][i], ocr_data['height'][i])
                                draw.rectangle([x, y, x + w, y + h], fill="black")

                # Save the edited image back to a byte stream
                img_byte_arr = io.BytesIO()
                pil_img.save(img_byte_arr, format=base_image["ext"])
                img_byte_arr = img_byte_arr.getvalue()

                # Replace the original image in the PDF with the redacted image
                image_rect = page.get_image_rects(xref)[0]
                page.insert_image(image_rect, stream=img_byte_arr)

        # Save the final redacted PDF
        doc.save(redacted_pdf_path)
        return output_filename

    except Exception as e:
        print(f"Error redacting PDF text: {e}")
        return None

def redact_image(image_path, output_filename):
    try:
        pil_img = Image.open(image_path)
        ocr_data = pytesseract.image_to_data(pil_img, output_type=pytesseract.Output.DICT)

        # Draw over sensitive text detected in the image
        draw = ImageDraw.Draw(pil_img)
        for i, word in enumerate(ocr_data['text']):
            if word.strip():  # Ensure word is not empty
                text_entities = ner_pipeline(word)
                for entity in text_entities:
                    if entity['entity_group'] in ["PER", "LOC", "ORG", "MISC"]:
                        # Get bounding box and draw a black rectangle over the sensitive word
                        x, y, w, h = (ocr_data['left'][i], ocr_data['top'][i],
                                      ocr_data['width'][i], ocr_data['height'][i])
                        draw.rectangle([x, y, x + w, y + h], fill="black")

        # Save the redacted image
        redacted_image_path = f"./redacted/{output_filename}"
        pil_img.save(redacted_image_path)

        return output_filename

    except Exception as e:
        print(f"Error redacting image text: {e}")
        return None

def redact_faces(image_path):
    # Load YOLO
    config_path = './yolo/yolov4.cfg'  # Path to config file
    weights_path = './yolo/yolov4.weights'  # Path to weights file

    # Check if files exist
    if not os.path.isfile(config_path) or not os.path.isfile(weights_path):
        raise FileNotFoundError("YOLO config or weights file not found.")

    # Create the network
    net = cv2.dnn.readNet(weights_path, config_path)

    # Load the image
    image = cv2.imread(image_path)
    height, width = image.shape[:2]

    # Prepare the image for the model
    blob = cv2.dnn.blobFromImage(image, 1/255.0, (416, 416), swapRB=True, crop=False)
    net.setInput(blob)

    # Get the output layer names
    layer_names = net.getLayerNames()
    output_layers = [layer_names[i - 1] for i in net.getUnconnectedOutLayers()]

    # Perform inference
    detections = net.forward(output_layers)

    # Loop through detected objects
    for detection in detections:
        for obj in detection:
            scores = obj[5:]  # Get the scores for the detected object
            class_id = np.argmax(scores)  # Get the class ID
            confidence = scores[class_id]  # Get the confidence score

            # Filter out weak detections and non-person classes
            if confidence > 0.5 and class_id == 0:  # 0 is usually 'person'
                center_x = int(obj[0] * width)
                center_y = int(obj[1] * height)
                w = int(obj[2] * width)
                h = int(obj[3] * height)

                # Get the bounding box coordinates
                x1 = int(center_x - w / 2)
                y1 = int(center_y - h / 2)
                x2 = int(center_x + w / 2)
                y2 = int(center_y + h / 2)

                # Redact the face by applying Gaussian blur
                face = image[y1:y2, x1:x2]
                blurred_face = cv2.GaussianBlur(face, (99, 99), 30)

                # Replace the face with the blurred face
                image[y1:y2, x1:x2] = blurred_face

    # Save the redacted image
    output_path = os.path.join('redacted', os.path.basename(image_path))
    cv2.imwrite(output_path, image)
    return os.path.basename(image_path)

def redact_metadata(file_path):

    filename, ext = os.path.splitext(os.path.basename(file_path))
    ext = ext.lower()

    if ext in ['.jpg', '.jpeg']:
        return remove_image_metadata(file_path, filename + '_redacted.jpg')
    elif ext == '.png':
        return remove_image_metadata(file_path, filename + '_redacted.png')
    elif ext == '.pdf':
        return redact_pdf_metadata(file_path, filename + '_redacted' + ext)
    elif ext in ['.docx', '.doc']:
        return redact_word_metadata(file_path, filename + '_redacted' + ext)
    
    return None  # Return None if the file type is not supported

def remove_image_metadata(image_path, output_filename):
    try:
        img = Image.open(image_path)

        # Remove EXIF metadata from images (JPEG, PNG)
        if hasattr(img, '_getexif'):  # Only JPEG format has EXIF
            exif_data = img.info.get('exif')
            if exif_data:
                img = img.convert('RGB')  # Convert to remove EXIF

        # Handle RGBA (transparency) images
        if img.mode == 'RGBA':
            img = img.convert('RGB')

        # Save without metadata
        redacted_image_path = os.path.join('./redacted', output_filename)
        img.save(redacted_image_path, quality=95)
        return output_filename
    except Exception as e:
        print(f"Error redacting image metadata: {e}")
        return None

def redact_pdf_metadata(pdf_path, output_filename):
    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()

        # Copy pages without metadata
        for page in reader.pages:
            writer.add_page(page)

        # Clear metadata if present
        writer.add_metadata({})

        redacted_pdf_path = os.path.join('./redacted', output_filename)
        with open(redacted_pdf_path, 'wb') as f:
            writer.write(f)

        return output_filename
    except Exception as e:
        print(f"Error redacting PDF metadata: {e}")
        return None

def redact_word_metadata(doc_path, output_filename):
    try:
        doc = Document(doc_path)

        # Remove document properties (author, title, etc.)
        core_props = doc.core_properties
        core_props.author = ""
        core_props.title = ""
        core_props.subject = ""
        core_props.keywords = ""
        core_props.comments = ""

        # Remove comments in document
        for comment in doc.element.xpath('//w:comment'):
            comment.getparent().remove(comment)

        # Remove track changes if enabled by accepting them
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.highlight_color:
                    run.clear()

        # Remove any custom XML properties and sensitive data
        custom_properties = doc.part
        if hasattr(custom_properties, 'remove_element'):
            custom_properties.remove_element(custom_properties)

        # Remove specific sensitive terms
        for paragraph in doc.paragraphs:
            if "sensitive" in paragraph.text.lower():
                paragraph.text = paragraph.text.replace("sensitive", "[REDACTED]")

        # Save the redacted document
        redacted_doc_path = os.path.join('./redacted', output_filename)
        doc.save(redacted_doc_path)
        return output_filename
    except Exception as e:
        print(f"Error redacting Word metadata: {e}")
        return None